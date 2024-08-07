import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from num2words import num2words
from datetime import datetime


def format_cpf(cpf):
    cpf = str(cpf).zfill(11)
    return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"

def format_date():
    today = datetime.now()
    months = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    day = today.day
    month = months[today.month - 1]
    year = today.year
    return f"{day:02d} de {month} de {year}"

def format_text(text, font_size):
    for run in text.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)


#Fill with your credentials before run the script
name = ""
marital_status = ""
profession = ""
profession_code = ""
profession_number = ""
rg = ""
cpf_number = ""
service = ""

#open excel file
file = "dados.xlsx"
df = pd.read_excel(file)

try:
    #save data from columns
    patient = df['paciente']
    cpf = df['cpf']
    value = df['valor']
    month = df['mes']
    year = df['ano']

    #start generating the receipts
    for i in range(len(patient)):
        formatted_cpf = format_cpf(cpf[i])
        formatted_value = f"{value[i]:.2f}".replace('.', ',')
        value_in_words = num2words(value[i], lang='pt_BR')

        title = "Recibo\n"
        text = (
            f"Eu, {name}, {marital_status}, {profession}, Inscrito no {profession_code} sob o Nº: {profession_number}, no RG Nº {rg} e no CPF: {cpf_number}, "
            f"declaro que recebi de {patient[i].capitalize()}, Inscrito no CPF: {formatted_cpf}, a importância de R$ {formatted_value} ({value_in_words} reais) "
            f"correspondente ao pagamento pelos atendimentos de {service} individual realizados com ele ao longo do mês de {month[i].capitalize()} de {year[i]} "
            f"({num2words(year[i], lang='pt_BR')})."
        )
        local_and_date = f"\n\nBelo Horizonte, {format_date()}"

        #start saving document
        doc = Document()

        #add and align title
        main_title = doc.add_paragraph(title)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_text(main_title, 16)
        
        #add and align paragraph
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        format_text(paragraph, 12)

        #add and align local and date
        local_and_date = doc.add_paragraph(local_and_date)
        local_and_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_text(local_and_date, 12)

        #add and align logo
        doc.add_picture('logo.png', width=Inches(4.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        filename = f"declaracao_{patient[i].capitalize()}.docx"
        
        if not os.path.exists('saved'):
            os.makedirs('saved')

        doc_path = os.path.join('saved', filename)
        doc.save(doc_path)

    print(f"\033[92mScript executado com sucesso! Foram salvos {len(patient)} arquivos na pasta 'receipt-generator/saved/'\033[0m")
except:
    print(f"\033[91mFalha ao abrir arquivo, verifique se o arquivo está na pasta atual e se o nome dele está escrito corretamente.\033[0m")
